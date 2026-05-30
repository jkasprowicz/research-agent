from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("/Users/joaokasprowicz/Documents/Projects/research-agent/projects/phd-positioning/outputs/plano_doutorado_ppgcc_submissao_final.md")
OUT = Path("/Users/joaokasprowicz/Documents/Projects/research-agent/projects/phd-positioning/outputs/plano_doutorado_ppgcc_submissao_final.docx")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = [
        ("Heading 1", 12, True, 6, 3),
        ("Heading 2", 12, True, 4, 2),
        ("Heading 3", 12, True, 2, 1),
    ]
    for name, size, bold, before, after in heading_specs:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Plano de Trabalho de Doutorado")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(12)
    r.bold = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run("• ").bold = True
    p.add_run(text)


def add_numbered(doc, num, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run(f"{num}. ").bold = True
    p.add_run(text)


def build():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_title(doc)

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            continue
        m = re.match(r"(\d+)\.\s+(.*)", line)
        if m:
            add_numbered(doc, m.group(1), m.group(2))
            continue
        p = doc.add_paragraph(style="Normal")
        p.add_run(line)

    doc.save(OUT)


if __name__ == "__main__":
    build()
