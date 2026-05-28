from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/joaokasprowicz/Documents/Projects/research-agent/projects/phd-positioning/outputs")
SRC = ROOT / "refined_four_page_plan_v2.md"
OUT = ROOT / "refined_four_page_plan_v2.docx"


def set_cell_margin(table, top=80, start=120, bottom=80, end=120):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = tcPr.first_child_found_in("w:tcMar")
            if tcMar is None:
                tcMar = OxmlElement("w:tcMar")
                tcPr.append(tcMar)
            for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
                node = tcMar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tcMar.append(node)
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 10),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("From Morphology Intelligence to Multimodal Clinical-Laboratory Intelligence in Hematology")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(19)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run("Doctoral proposal refinement for PPGCC/UFSC selection")
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2.font.size = Pt(11)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_metadata_table(doc):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.75), Inches(4.75)]
    content = [
        ("Positioning axis", "Trustworthy multimodal clinical-laboratory AI in hematology"),
        ("Core thesis scope", "Peripheral blood smear morphology + CBC parameters + analyzer flags + structured laboratory context"),
    ]
    for i, (label, value) in enumerate(content):
        for j, text in enumerate((label, value)):
            cell = table.cell(i, j)
            cell.width = widths[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            run.font.size = Pt(10.5)
            if j == 0:
                run.bold = True
    set_cell_margin(table)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(u"\u2022 ")
    run.bold = True
    run2 = p.add_run(text)
    run2.font.name = "Calibri"
    run2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_numbered(doc, idx, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(f"{idx}. ")
    run.bold = True
    p.add_run(text)


def parse_markdown(doc, text):
    lines = text.splitlines()
    numbered_buffer = None
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            continue
        m = re.match(r"(\d+)\.\s+(.*)", line)
        if m:
            add_numbered(doc, m.group(1), m.group(2))
            continue
        p = doc.add_paragraph(style="Normal")
        p.add_run(line)


def build():
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc)
    base_styles(doc)
    add_title_block(doc)
    add_metadata_table(doc)
    parse_markdown(doc, text)
    doc.save(OUT)


if __name__ == "__main__":
    build()
